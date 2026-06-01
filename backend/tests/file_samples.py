from io import BytesIO

from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter
from pypdf.generic import DictionaryObject, NameObject, StreamObject


def make_text_pdf(text: str) -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
    )
    stream = StreamObject()
    stream._data = f"BT /F1 12 Tf 72 720 Td ({_escape_pdf_text(text)}) Tj ET".encode("ascii")
    page[NameObject("/Contents")] = writer._add_object(stream)

    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def make_blank_pdf() -> bytes:
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)

    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def make_docx(paragraph_text: str, table_rows: list[list[str]] | None = None) -> bytes:
    document = Document()
    document.add_paragraph(paragraph_text)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=max(len(row) for row in table_rows))
        for row_index, row in enumerate(table_rows):
            for col_index, value in enumerate(row):
                table.cell(row_index, col_index).text = value

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_xlsx(sheet_name: str, rows: list[list[str | int]]) -> bytes:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = sheet_name
    for row in rows:
        worksheet.append(row)

    buffer = BytesIO()
    workbook.save(buffer)
    workbook.close()
    return buffer.getvalue()


def _escape_pdf_text(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
