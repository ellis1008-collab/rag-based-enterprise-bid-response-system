from pathlib import Path
import csv
import io

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import load_workbook
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.models import AgentRun, BidResponse


def test_generate_responses_for_sample_project(client: TestClient, db_session: Session) -> None:
    project = client.post(
        "/api/rfp/projects",
        json={"name": "响应矩阵项目", "customer_name": "样例客户"},
    ).json()
    project_id = project["id"]
    upload_sample_rfp(client, project_id)
    upload_sample_product_docs(client)

    requirements = client.post(f"/api/rfp/projects/{project_id}/extract-requirements").json()
    response = client.post(f"/api/rfp/projects/{project_id}/generate-responses")

    assert response.status_code == 200
    responses = response.json()
    assert len(responses) == len(requirements)
    assert any(item["risk_level"] == "medium" for item in responses)
    assert all(item["source_chunks"] for item in responses)
    assert all(item["human_status"] == "pending" for item in responses)
    assert all(item["human_note"] == "" for item in responses)
    assert all(item["updated_at"] for item in responses)

    list_response = client.get(f"/api/rfp/projects/{project_id}/responses")
    assert list_response.status_code == 200
    assert len(list_response.json()) == len(requirements)

    by_requirement_id = {item["requirement_id"]: item for item in responses}
    requirement_by_code = {item["requirement_code"]: item for item in requirements}
    expected = {
        "REQ-001": ("satisfied", "low"),
        "REQ-002": ("satisfied", "low"),
        "REQ-003": ("satisfied", "low"),
        "REQ-004": ("satisfied", "low"),
        "REQ-005": ("satisfied", "low"),
        "REQ-006": ("satisfied", "low"),
        "REQ-007": ("partial", "medium"),
        "REQ-008": ("partial", "medium"),
    }
    for requirement_code, (match_status, risk_level) in expected.items():
        requirement = requirement_by_code[requirement_code]
        item = by_requirement_id[requirement["id"]]
        assert item["match_status"] == match_status
        assert item["risk_level"] == risk_level

    agent_run = db_session.scalar(
        select(AgentRun)
        .where(AgentRun.project_id == project_id, AgentRun.run_type == "generate_responses")
        .order_by(AgentRun.id.desc())
    )
    assert agent_run is not None
    assert agent_run.status == "succeeded"
    assert agent_run.finished_at is not None
    assert agent_run.steps_json["requirement_count"] == len(requirements)
    assert agent_run.steps_json["generated_response_count"] == len(responses)
    assert agent_run.steps_json["risk_summary"]["medium"] >= 1


def test_generate_responses_agent_run_steps_include_retrieval_details(
    client: TestClient,
    db_session: Session,
) -> None:
    project_id = create_sample_response_matrix(client)

    agent_run = db_session.scalar(
        select(AgentRun)
        .where(AgentRun.project_id == project_id, AgentRun.run_type == "generate_responses")
        .order_by(AgentRun.id.desc())
    )
    assert agent_run is not None
    steps = agent_run.steps_json["steps"]
    assert [step["name"] for step in steps] == [
        "load_requirements",
        "retrieve_knowledge",
        "call_llm_for_each_requirement",
        "validate_schema",
        "save_bid_responses",
        "build_risk_summary",
    ]
    assert all(step["status"] == "completed" for step in steps)
    assert steps[0]["requirement_count"] == 8
    assert steps[1]["retrieved_chunk_count"] >= 8
    assert len(steps[1]["retrievals"]) == 8
    assert steps[1]["retrievals"][0]["query"]
    assert steps[1]["retrievals"][0]["retrieved_chunks"]
    assert "content_summary" in steps[1]["retrievals"][0]["retrieved_chunks"][0]
    assert steps[1]["retrievals"][0]["retrieved_chunks"][0]["retriever_type"] == "simple"
    assert steps[2]["prompt_type"] == "generate_response"
    assert steps[3]["schema"] == "BidResponseGenerationResult"
    assert steps[5]["risk_summary"]["medium"] == 2


def test_generate_responses_replaces_existing_responses(client: TestClient, db_session: Session) -> None:
    project = client.post(
        "/api/rfp/projects",
        json={"name": "重复响应矩阵项目", "customer_name": "样例客户"},
    ).json()
    project_id = project["id"]
    upload_sample_rfp(client, project_id)
    upload_sample_product_docs(client)
    client.post(f"/api/rfp/projects/{project_id}/extract-requirements")

    first = client.post(f"/api/rfp/projects/{project_id}/generate-responses").json()
    second = client.post(f"/api/rfp/projects/{project_id}/generate-responses").json()

    assert len(first) == len(second)
    stored_count = len(list(db_session.scalars(select(BidResponse).where(BidResponse.project_id == project_id))))
    assert stored_count == len(second)


def test_update_response_human_review_fields(client: TestClient) -> None:
    project_id = create_sample_response_matrix(client)
    response_item = client.get(f"/api/rfp/projects/{project_id}/responses").json()[0]

    patch_response = client.patch(
        f"/api/rfp/projects/{project_id}/responses/{response_item['id']}",
        json={
            "human_status": "confirmed",
            "human_note": "售前已复核",
            "response_text": "人工修订后的响应说明",
        },
    )

    assert patch_response.status_code == 200
    updated = patch_response.json()
    assert updated["human_status"] == "confirmed"
    assert updated["human_note"] == "售前已复核"
    assert updated["response_text"] == "人工修订后的响应说明"
    assert updated["updated_at"] >= response_item["updated_at"]

    refreshed = client.get(f"/api/rfp/projects/{project_id}/responses").json()[0]
    assert refreshed["human_status"] == "confirmed"
    assert refreshed["human_note"] == "售前已复核"
    assert refreshed["response_text"] == "人工修订后的响应说明"


def test_update_response_rejects_response_from_other_project(client: TestClient) -> None:
    project_id = create_sample_response_matrix(client)
    other_project = client.post(
        "/api/rfp/projects",
        json={"name": "其他项目", "customer_name": "样例客户"},
    ).json()
    response_item = client.get(f"/api/rfp/projects/{project_id}/responses").json()[0]

    patch_response = client.patch(
        f"/api/rfp/projects/{other_project['id']}/responses/{response_item['id']}",
        json={"human_status": "rejected"},
    )

    assert patch_response.status_code == 404


def test_update_response_rejects_invalid_review_status(client: TestClient) -> None:
    project_id = create_sample_response_matrix(client)
    response_item = client.get(f"/api/rfp/projects/{project_id}/responses").json()[0]

    patch_response = client.patch(
        f"/api/rfp/projects/{project_id}/responses/{response_item['id']}",
        json={"human_status": "done"},
    )

    assert patch_response.status_code == 422


def test_risk_report_counts_human_review_statuses(client: TestClient) -> None:
    project_id = create_sample_response_matrix(client)
    responses = client.get(f"/api/rfp/projects/{project_id}/responses").json()

    client.patch(
        f"/api/rfp/projects/{project_id}/responses/{responses[0]['id']}",
        json={"human_status": "confirmed"},
    )
    client.patch(
        f"/api/rfp/projects/{project_id}/responses/{responses[1]['id']}",
        json={"human_status": "rejected"},
    )

    report = client.get(f"/api/rfp/projects/{project_id}/risk-report").json()
    assert report["pending_review_count"] == 6
    assert report["confirmed_count"] == 1
    assert report["rejected_count"] == 1


def test_export_responses_csv_and_risk_report(client: TestClient) -> None:
    project_id = create_sample_response_matrix(client)
    response_item = client.get(f"/api/rfp/projects/{project_id}/responses").json()[0]
    client.patch(
        f"/api/rfp/projects/{project_id}/responses/{response_item['id']}",
        json={"human_status": "confirmed", "human_note": "CSV 备注"},
    )

    csv_response = client.get(f"/api/rfp/projects/{project_id}/responses/export-csv")
    assert csv_response.status_code == 200
    assert "bidpilot_" in csv_response.headers["content-disposition"]
    assert str(project_id) in csv_response.headers["content-disposition"]

    csv_text = csv_response.content.decode("utf-8-sig")
    rows = list(csv.DictReader(io.StringIO(csv_text)))
    assert len(rows) == 8
    assert set(rows[0]) == {
        "requirement_code",
        "category",
        "priority",
        "requirement_content",
        "match_status",
        "risk_level",
        "response_text",
        "source_summary",
        "human_status",
        "human_note",
    }
    assert rows[0]["requirement_code"] == "REQ-001"
    assert rows[0]["source_summary"]
    assert rows[0]["human_status"] == "confirmed"
    assert rows[0]["human_note"] == "CSV 备注"

    report_response = client.get(f"/api/rfp/projects/{project_id}/risk-report")
    assert report_response.status_code == 200
    report = report_response.json()
    assert report["total_requirements"] == 8
    assert report["satisfied_count"] == 6
    assert report["partial_count"] == 2
    assert report["unsupported_count"] == 0
    assert report["low_risk_count"] == 6
    assert report["medium_risk_count"] == 2
    assert report["high_risk_count"] == 0
    assert report["pending_review_count"] == 7
    assert report["confirmed_count"] == 1
    assert report["rejected_count"] == 0
    assert len(report["risk_items"]) == 2
    assert report["pending_confirmation_items"] == []


def test_export_responses_xlsx_contains_review_fields(client: TestClient) -> None:
    project_id = create_sample_response_matrix(client)
    response_item = client.get(f"/api/rfp/projects/{project_id}/responses").json()[0]
    client.patch(
        f"/api/rfp/projects/{project_id}/responses/{response_item['id']}",
        json={"human_status": "confirmed", "human_note": "Excel 备注"},
    )

    xlsx_response = client.get(f"/api/rfp/projects/{project_id}/responses/export-xlsx")

    assert xlsx_response.status_code == 200
    assert xlsx_response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    assert "bidpilot_" in xlsx_response.headers["content-disposition"]
    workbook = load_workbook(io.BytesIO(xlsx_response.content))
    worksheet = workbook["技术响应矩阵"]
    headers = [cell.value for cell in worksheet[1]]
    assert headers == [
        "需求编号",
        "分类",
        "优先级",
        "客户要求",
        "是否满足",
        "风险等级",
        "技术响应说明",
        "引用来源摘要",
        "人工复核状态",
        "人工备注",
    ]
    values = [cell.value for cell in worksheet[2]]
    assert values[0] == "REQ-001"
    assert values[8] == "已确认"
    assert values[9] == "Excel 备注"
    assert worksheet.max_row == 9


def test_export_proposal_docx_contains_project_and_summary(client: TestClient) -> None:
    project_id = create_sample_response_matrix(client)

    docx_response = client.get(f"/api/rfp/projects/{project_id}/proposal/export-docx")

    assert docx_response.status_code == 200
    assert docx_response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "bidpilot_" in docx_response.headers["content-disposition"]
    document = Document(io.BytesIO(docx_response.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    full_text = f"{text}\n{table_text}"
    assert "导出测试项目 投标响应初稿" in full_text
    assert "客户名称：样例客户" in full_text
    assert "中高风险数量" in full_text
    assert "待人工确认数量" in full_text
    assert "本文档为 AI 生成初稿" in full_text


def test_export_deliverables_return_clear_error_without_response_matrix(client: TestClient) -> None:
    project = client.post(
        "/api/rfp/projects",
        json={"name": "空矩阵项目", "customer_name": "样例客户"},
    ).json()
    project_id = project["id"]

    xlsx_response = client.get(f"/api/rfp/projects/{project_id}/responses/export-xlsx")
    docx_response = client.get(f"/api/rfp/projects/{project_id}/proposal/export-docx")

    assert xlsx_response.status_code == 400
    assert xlsx_response.json()["detail"] == "No response matrix available for export."
    assert docx_response.status_code == 400
    assert docx_response.json()["detail"] == "No response matrix available for export."


def create_sample_response_matrix(client: TestClient) -> int:
    project = client.post(
        "/api/rfp/projects",
        json={"name": "导出测试项目", "customer_name": "样例客户"},
    ).json()
    project_id = project["id"]
    upload_sample_rfp(client, project_id)
    upload_sample_product_docs(client)
    client.post(f"/api/rfp/projects/{project_id}/extract-requirements")
    client.post(f"/api/rfp/projects/{project_id}/generate-responses")
    return project_id


def upload_sample_rfp(client: TestClient, project_id: int) -> None:
    sample_path = Path(__file__).resolve().parents[2] / "sample-data" / "sample_rfp.txt"
    content = sample_path.read_text(encoding="utf-8")
    response = client.post(
        f"/api/rfp/projects/{project_id}/documents/upload",
        files={"file": ("sample_rfp.txt", content.encode("utf-8"), "text/plain")},
    )
    assert response.status_code == 201


def upload_sample_product_docs(client: TestClient) -> None:
    sample_path = Path(__file__).resolve().parents[2] / "sample-data" / "product_docs.txt"
    content = sample_path.read_text(encoding="utf-8")
    response = client.post(
        "/api/knowledge/upload",
        files={"file": ("product_docs.txt", content.encode("utf-8"), "text/plain")},
    )
    assert response.status_code == 201
