import io
from dataclasses import dataclass

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from sqlalchemy.orm import Session

from app.models import RfpProject
from app.models.timestamps import utc_now
from app.services.bid_response_service import ResponseExportRow, build_response_export_rows, build_risk_report


class DeliverableExportError(Exception):
    pass


@dataclass(frozen=True)
class BinaryExport:
    content: bytes


XLSX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

EXCEL_HEADERS = [
    "需求编号",
    "分类",
    "优先级",
    "客户要求",
    "是否满足",
    "风险等级",
    "技术响应说明",
    "引用来源摘要",
    "人工复核状态",
    "人工备注",
]

MATCH_LABELS = {
    "satisfied": "满足",
    "partial": "部分满足",
    "unsupported": "不支持",
}
RISK_LABELS = {
    "low": "低风险",
    "medium": "中风险",
    "high": "高风险",
}
HUMAN_STATUS_LABELS = {
    "pending": "待确认",
    "confirmed": "已确认",
    "rejected": "已驳回",
}


def build_response_matrix_xlsx(db: Session, project_id: int) -> BinaryExport:
    rows = _response_rows_or_error(db, project_id)
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "技术响应矩阵"
    worksheet.append(EXCEL_HEADERS)

    for row in rows:
        worksheet.append(
            [
                row.requirement_code,
                row.category,
                row.priority,
                row.requirement_content,
                _match_label(row.match_status),
                _risk_label(row.risk_level),
                row.response_text,
                row.source_summary,
                _human_status_label(row.human_status),
                row.human_note,
            ]
        )

    _style_response_matrix_sheet(worksheet, rows)
    output = io.BytesIO()
    workbook.save(output)
    return BinaryExport(content=output.getvalue())


def build_proposal_docx(db: Session, project: RfpProject) -> BinaryExport:
    rows = _response_rows_or_error(db, project.id)
    report = build_risk_report(db, project.id)

    document = Document()
    _configure_document_styles(document)
    _add_title(document, project)
    _add_summary_section(document, report)
    _add_matrix_section(document, rows)
    _add_risk_section(document, rows)
    _add_note_section(document)

    output = io.BytesIO()
    document.save(output)
    return BinaryExport(content=output.getvalue())


def _response_rows_or_error(db: Session, project_id: int) -> list[ResponseExportRow]:
    rows = build_response_export_rows(db, project_id)
    if not rows:
        raise DeliverableExportError("No response matrix available for export.")
    return rows


def _style_response_matrix_sheet(worksheet, rows: list[ResponseExportRow]) -> None:
    header_fill = PatternFill("solid", fgColor="0F766E")
    header_font = Font(bold=True, color="FFFFFF")
    thin_border = Border(
        left=Side(style="thin", color="CBD5E1"),
        right=Side(style="thin", color="CBD5E1"),
        top=Side(style="thin", color="CBD5E1"),
        bottom=Side(style="thin", color="CBD5E1"),
    )
    medium_fill = PatternFill("solid", fgColor="FEF3C7")
    high_fill = PatternFill("solid", fgColor="FEE2E2")
    status_fills = {
        "pending": PatternFill("solid", fgColor="FEF3C7"),
        "confirmed": PatternFill("solid", fgColor="D1FAE5"),
        "rejected": PatternFill("solid", fgColor="FEE2E2"),
    }

    for cell in worksheet[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin_border

    for row_index, row in enumerate(rows, start=2):
        row_fill = None
        if row.risk_level == "high":
            row_fill = high_fill
        elif row.risk_level == "medium":
            row_fill = medium_fill

        for cell in worksheet[row_index]:
            if row_fill is not None:
                cell.fill = row_fill
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = thin_border

        status_cell = worksheet.cell(row=row_index, column=9)
        status_cell.fill = status_fills.get(row.human_status or "pending", status_fills["pending"])
        status_cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        risk_cell = worksheet.cell(row=row_index, column=6)
        risk_cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    widths = [14, 14, 10, 36, 12, 12, 46, 42, 14, 28]
    for index, width in enumerate(widths, start=1):
        worksheet.column_dimensions[get_column_letter(index)].width = width
    worksheet.freeze_panes = "A2"
    worksheet.auto_filter.ref = worksheet.dimensions


def _configure_document_styles(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.bold = True


def _add_title(document: Document, project: RfpProject) -> None:
    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(f"{project.name} 投标响应初稿")
    title_run.bold = True
    title_run.font.size = Pt(18)

    document.add_paragraph(f"客户名称：{project.customer_name}")
    document.add_paragraph(f"生成时间：{utc_now().strftime('%Y-%m-%d %H:%M:%S UTC')}")


def _add_summary_section(document: Document, report) -> None:
    document.add_heading("一、响应摘要", level=1)
    high_or_medium_risk_count = report.medium_risk_count + report.high_risk_count
    summary_rows = [
        ("需求总数", report.total_requirements),
        ("满足数量", report.satisfied_count),
        ("部分满足数量", report.partial_count),
        ("不支持数量", report.unsupported_count),
        ("中高风险数量", high_or_medium_risk_count),
        ("待人工确认数量", report.pending_review_count),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.rows[0].cells[0].text = "指标"
    table.rows[0].cells[1].text = "数量"
    for label, value in summary_rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)
    _style_docx_table(table)


def _add_matrix_section(document: Document, rows: list[ResponseExportRow]) -> None:
    document.add_heading("二、技术响应矩阵", level=1)
    for row in rows:
        document.add_heading(row.requirement_code, level=2)
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _add_key_value_row(table, "客户要求", row.requirement_content)
        _add_key_value_row(table, "是否满足", _match_label(row.match_status))
        _add_key_value_row(table, "技术响应说明", row.response_text)
        _add_key_value_row(table, "风险等级", _risk_label(row.risk_level))
        _add_key_value_row(table, "引用来源", row.source_summary or "无")
        _add_key_value_row(table, "人工复核状态", _human_status_label(row.human_status))
        _add_key_value_row(table, "人工备注", row.human_note or "无")
        _style_docx_table(table)


def _add_risk_section(document: Document, rows: list[ResponseExportRow]) -> None:
    document.add_heading("三、风险与待确认事项", level=1)
    flagged_rows = [
        row
        for row in rows
        if row.match_status in {"partial", "unsupported"}
        or row.risk_level in {"medium", "high"}
        or (row.human_status or "pending") == "pending"
    ]
    if not flagged_rows:
        document.add_paragraph("暂无风险或待确认事项。")
        return

    for row in flagged_rows:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(
            f"{row.requirement_code}：{_match_label(row.match_status)} / "
            f"{_risk_label(row.risk_level)} / {_human_status_label(row.human_status)}"
        )
        if row.risk_level in {"medium", "high"} or row.match_status != "satisfied":
            run.bold = True
            run.font.color.rgb = RGBColor(180, 83, 9)
        document.add_paragraph(row.requirement_content)


def _add_note_section(document: Document) -> None:
    document.add_heading("四、说明", level=1)
    document.add_paragraph("本文档为 AI 生成初稿，需售前或解决方案团队最终确认后方可作为正式投标文件使用。")


def _add_key_value_row(table, key: str, value: str) -> None:
    cells = table.add_row().cells
    cells[0].text = key
    cells[1].text = value


def _style_docx_table(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    for cell in table.columns[0].cells:
        cell.width = Inches(1.4)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _match_label(value: str) -> str:
    return MATCH_LABELS.get(value, value)


def _risk_label(value: str) -> str:
    return RISK_LABELS.get(value, value)


def _human_status_label(value: str) -> str:
    return HUMAN_STATUS_LABELS.get(value or "pending", value or "pending")
