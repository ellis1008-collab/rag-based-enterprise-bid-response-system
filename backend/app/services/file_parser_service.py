from dataclasses import dataclass
from io import BytesIO
import logging
from pathlib import Path
from typing import Any

from fastapi import UploadFile


SUPPORTED_FILE_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".xlsx"}
TEXT_EXTENSIONS = {".txt", ".md"}
PDF_OCR_RENDER_SCALE = 2.0
OCR_MIN_CONFIDENCE = 0.5

logger = logging.getLogger(__name__)


class FileParserError(ValueError):
    pass


@dataclass(frozen=True)
class ParsedFile:
    filename: str
    content_text: str
    extension: str


class FileParserService:
    def __init__(self, supported_extensions: set[str] | None = None) -> None:
        self.supported_extensions = supported_extensions or SUPPORTED_FILE_EXTENSIONS

    async def parse_upload_file(self, file: UploadFile) -> ParsedFile:
        filename = file.filename or ""
        raw_content = await file.read()
        return self.parse_bytes(filename=filename, raw_content=raw_content)

    def parse_bytes(self, filename: str, raw_content: bytes) -> ParsedFile:
        extension = Path(filename).suffix.lower()
        if extension not in self.supported_extensions:
            supported = ", ".join(sorted(self.supported_extensions))
            raise FileParserError(f"Unsupported file format. Supported formats: {supported}.")

        if not raw_content:
            raise FileParserError("Uploaded file is empty.")

        if extension in TEXT_EXTENSIONS:
            content_text = self._parse_text(raw_content)
        elif extension == ".pdf":
            content_text = self._parse_pdf(raw_content)
        elif extension == ".docx":
            content_text = self._parse_docx(raw_content)
        elif extension == ".xlsx":
            content_text = self._parse_xlsx(raw_content)
        else:
            raise FileParserError(f"Unsupported file format: {extension}.")

        content_text = self._normalize_text(content_text)
        if not content_text.strip():
            raise FileParserError("Parsed file contains no readable text.")

        return ParsedFile(filename=filename, content_text=content_text, extension=extension)

    def _parse_text(self, raw_content: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
            try:
                return raw_content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return raw_content.decode("utf-8", errors="replace")

    def _parse_pdf(self, raw_content: bytes) -> str:
        content_text = self._parse_pdf_text_layer(raw_content)
        if content_text.strip():
            return content_text

        logger.info("PDF text layer is empty; falling back to OCR parsing.")
        return self._parse_pdf_with_ocr(raw_content)

    def _parse_pdf_text_layer(self, raw_content: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise FileParserError("PDF parsing dependency is not installed.") from exc

        try:
            reader = PdfReader(BytesIO(raw_content))
            if reader.is_encrypted:
                reader.decrypt("")
            page_texts = [page.extract_text() or "" for page in reader.pages]
        except Exception as exc:
            raise FileParserError(
                "Failed to parse PDF file."
            ) from exc

        return "\n\n".join(text.strip() for text in page_texts if text.strip())

    def _parse_pdf_with_ocr(self, raw_content: bytes) -> str:
        try:
            import pypdfium2 as pdfium
            from rapidocr import RapidOCR
        except ImportError as exc:
            raise FileParserError(
                "PDF file contains no embedded readable text, and OCR dependencies are not installed."
            ) from exc

        document = None
        try:
            document = pdfium.PdfDocument(raw_content)
            ocr = RapidOCR()
            page_texts: list[str] = []
            for page_index in range(len(document)):
                page = document[page_index]
                image = None
                try:
                    image = page.render(scale=PDF_OCR_RENDER_SCALE).to_pil()
                    page_text = self._ocr_image_text(ocr, image)
                    if page_text:
                        page_texts.append(page_text)
                except Exception:
                    logger.warning("OCR parsing failed for PDF page %s.", page_index + 1, exc_info=True)
                finally:
                    if image is not None:
                        image.close()
                    page.close()
        except FileParserError:
            raise
        except Exception as exc:
            raise FileParserError("Failed to parse image-based PDF file with OCR.") from exc
        finally:
            if document is not None:
                document.close()

        return "\n\n".join(page_texts)

    def _ocr_image_text(self, ocr: Any, image: Any) -> str:
        result = ocr(image)
        texts = getattr(result, "txts", None)
        scores = getattr(result, "scores", None)
        if texts is None and isinstance(result, tuple) and result:
            texts = [line[1] for line in result[0] or [] if len(line) >= 2]
            scores = [line[2] for line in result[0] or [] if len(line) >= 3]
        if not texts:
            return ""

        if not scores:
            return "\n".join(str(text).strip() for text in texts if str(text).strip())

        lines: list[str] = []
        for text, score in zip(texts, scores, strict=False):
            if float(score) < OCR_MIN_CONFIDENCE:
                continue
            line = str(text).strip()
            if line:
                lines.append(line)
        return "\n".join(lines)

    def _parse_docx(self, raw_content: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise FileParserError("DOCX parsing dependency is not installed.") from exc

        try:
            document = Document(BytesIO(raw_content))
            sections: list[str] = []
            sections.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    row_text = "\t".join(cell for cell in cells if cell)
                    if row_text:
                        sections.append(row_text)
        except Exception as exc:
            raise FileParserError("Failed to parse DOCX file.") from exc

        return "\n".join(sections)

    def _parse_xlsx(self, raw_content: bytes) -> str:
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise FileParserError("XLSX parsing dependency is not installed.") from exc

        try:
            workbook = load_workbook(BytesIO(raw_content), read_only=True, data_only=True)
            sections: list[str] = []
            for worksheet in workbook.worksheets:
                rows = self._worksheet_rows(worksheet)
                if rows:
                    sections.append(f"Sheet: {worksheet.title}")
                    sections.extend(rows)
            workbook.close()
        except Exception as exc:
            raise FileParserError("Failed to parse XLSX file.") from exc

        return "\n".join(sections)

    def _worksheet_rows(self, worksheet: Any) -> list[str]:
        rows: list[str] = []
        for row in worksheet.iter_rows(values_only=True):
            cells = [self._cell_to_text(value) for value in row]
            row_text = "\t".join(cell for cell in cells if cell)
            if row_text:
                rows.append(row_text)
        return rows

    def _cell_to_text(self, value: Any) -> str:
        if value is None:
            return ""
        return str(value).strip()

    def _normalize_text(self, text: str) -> str:
        return text.replace("\r\n", "\n").replace("\r", "\n")
